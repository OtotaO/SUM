#!/usr/bin/env python3
"""
multimodal_processor.py - Multi-Modal Content Processing Engine

This module provides comprehensive multi-modal processing capabilities for the SUM platform,
supporting text, PDF, DOCX, images, and other document formats with unified processing pipeline.

Features:
- PDF text extraction with layout preservation
- DOCX document processing with formatting awareness
- OCR for image-based text extraction
- Content type detection and routing
- Unified processing pipeline for all formats
- Local model integration with Ollama
- Vision-language model support

Author: ototao
License: Apache License 2.0
"""

import os
import logging
import mimetypes
from typing import Dict, List, Any, Optional, Tuple, Union
from dataclasses import dataclass
from pathlib import Path
import hashlib
import time
from enum import Enum

# Core document processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PDF processing not available. Install: pip install PyPDF2 pdfplumber")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("DOCX processing not available. Install: pip install python-docx")

# OCR capabilities
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR not available. Install: pip install Pillow pytesseract")

# Vision-language models
try:
    import requests
    import base64
    VISION_AVAILABLE = True
except ImportError:
    VISION_AVAILABLE = False
    logging.warning("Vision processing not available. Install: pip install requests")

# Import SUM components
from SUM import HierarchicalDensificationEngine

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ContentType(Enum):
    """Supported content types for multi-modal processing."""
    TEXT = "text"
    PDF = "pdf"
    DOCX = "docx"
    IMAGE = "image"
    HTML = "html"
    MARKDOWN = "markdown"
    UNKNOWN = "unknown"


@dataclass
class ProcessingResult:
    """Result of multi-modal processing."""
    content_type: ContentType
    extracted_text: str
    metadata: Dict[str, Any]
    processing_time: float
    confidence_score: float
    error_message: Optional[str] = None


@dataclass
class DocumentMetadata:
    """Metadata extracted from documents."""
    title: Optional[str] = None
    author: Optional[str] = None
    creation_date: Optional[str] = None
    modification_date: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    language: Optional[str] = None
    file_size: Optional[int] = None


class MultiModalProcessor:
    """
    Advanced multi-modal content processor with local model integration.
    
    This class handles various content types and provides unified processing
    with support for local AI models through Ollama integration.
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """Initialize the multi-modal processor."""
        self.config = config or {}
        self.hierarchical_engine = HierarchicalDensificationEngine()
        self.ollama_client = self._init_ollama()
        self._supported_formats = self._get_supported_formats()
        
        logger.info(f"MultiModalProcessor initialized with {len(self._supported_formats)} formats")
    
    def _init_ollama(self):
        """Initialize Ollama client for local model support."""
        try:
            import ollama
            client = ollama.Client()
            # Test connection
            models = client.list()
            logger.info(f"Ollama initialized with {len(models.get('models', []))} models")
            return client
        except ImportError:
            logger.warning("Ollama not available. Install: pip install ollama")
            return None
        except Exception as e:
            logger.warning(f"Ollama connection failed: {e}")
            return None
    
    def _get_supported_formats(self) -> Dict[str, ContentType]:
        """Get supported file formats and their content types."""
        formats = {
            '.txt': ContentType.TEXT,
            '.md': ContentType.MARKDOWN,
            '.html': ContentType.HTML,
            '.htm': ContentType.HTML,
        }
        
        if PDF_AVAILABLE:
            formats['.pdf'] = ContentType.PDF
            
        if DOCX_AVAILABLE:
            formats['.docx'] = ContentType.DOCX
            formats['.doc'] = ContentType.DOCX
            
        if OCR_AVAILABLE:
            formats.update({
                '.png': ContentType.IMAGE,
                '.jpg': ContentType.IMAGE,
                '.jpeg': ContentType.IMAGE,
                '.gif': ContentType.IMAGE,
                '.bmp': ContentType.IMAGE,
                '.tiff': ContentType.IMAGE,
            })
        
        return formats
    
    def detect_content_type(self, file_path: str) -> ContentType:
        """Detect content type from file path and MIME type."""
        try:
            path = Path(file_path)
            extension = path.suffix.lower()
            
            # Check by extension first
            if extension in self._supported_formats:
                return self._supported_formats[extension]
            
            # Fall back to MIME type detection
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type:
                if mime_type.startswith('text/'):
                    return ContentType.TEXT
                elif mime_type == 'application/pdf':
                    return ContentType.PDF
                elif mime_type.startswith('image/'):
                    return ContentType.IMAGE
                elif 'officedocument' in mime_type:
                    return ContentType.DOCX
            
            return ContentType.UNKNOWN
            
        except Exception as e:
            logger.error(f"Error detecting content type: {e}")
            return ContentType.UNKNOWN
    
    def process_file(self, file_path: str, **kwargs) -> ProcessingResult:
        """Process a file using appropriate handler based on content type."""
        start_time = time.time()
        
        try:
            # Detect content type
            content_type = self.detect_content_type(file_path)
            
            # Route to appropriate processor
            if content_type == ContentType.PDF:
                result = self._process_pdf(file_path, **kwargs)
            elif content_type == ContentType.DOCX:
                result = self._process_docx(file_path, **kwargs)
            elif content_type == ContentType.IMAGE:
                result = self._process_image(file_path, **kwargs)
            elif content_type == ContentType.HTML:
                result = self._process_html(file_path, **kwargs)
            elif content_type == ContentType.MARKDOWN:
                result = self._process_markdown(file_path, **kwargs)
            elif content_type == ContentType.TEXT:
                result = self._process_text(file_path, **kwargs)
            else:
                return ProcessingResult(
                    content_type=content_type,
                    extracted_text="",
                    metadata={},
                    processing_time=time.time() - start_time,
                    confidence_score=0.0,
                    error_message=f"Unsupported content type: {content_type}"
                )
            
            # Add processing time
            result.processing_time = time.time() - start_time
            
            # Enhance with hierarchical processing if text extracted
            if result.extracted_text and len(result.extracted_text.strip()) > 50:
                result = self._enhance_with_hierarchical_processing(result, **kwargs)
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return ProcessingResult(
                content_type=ContentType.UNKNOWN,
                extracted_text="",
                metadata={},
                processing_time=time.time() - start_time,
                confidence_score=0.0,
                error_message=str(e)
            )
    
    def _process_pdf(self, file_path: str, **kwargs) -> ProcessingResult:
        """Process PDF files with advanced text extraction."""
        if not PDF_AVAILABLE:
            return ProcessingResult(
                content_type=ContentType.PDF,
                extracted_text="",
                metadata={},
                processing_time=0.0,
                confidence_score=0.0,
                error_message="PDF processing not available"
            )
        
        try:
            text_content = []
            metadata = DocumentMetadata()
            
            # Try pdfplumber first for better text extraction
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    metadata.page_count = len(pdf.pages)
                    
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(page_text)
                    
                    # Extract metadata
                    if pdf.metadata:
                        metadata.title = pdf.metadata.get('Title')
                        metadata.author = pdf.metadata.get('Author')
                        metadata.creation_date = str(pdf.metadata.get('CreationDate', ''))
                        
            except ImportError:
                # Fall back to PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata.page_count = len(pdf_reader.pages)
                    
                    for page in pdf_reader.pages:
                        text_content.append(page.extract_text())
                    
                    # Extract metadata
                    if pdf_reader.metadata:
                        metadata.title = pdf_reader.metadata.get('/Title')
                        metadata.author = pdf_reader.metadata.get('/Author')
            
            extracted_text = '\n\n'.join(text_content)
            metadata.word_count = len(extracted_text.split())
            metadata.file_size = os.path.getsize(file_path)
            
            return ProcessingResult(
                content_type=ContentType.PDF,
                extracted_text=extracted_text,
                metadata=metadata.__dict__,
                processing_time=0.0,  # Will be set by caller
                confidence_score=0.9 if extracted_text.strip() else 0.3
            )
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return ProcessingResult(
                content_type=ContentType.PDF,
                extracted_text="",
                metadata={},
                processing_time=0.0,
                confidence_score=0.0,
                error_message=str(e)
            )
    
    def _process_docx(self, file_path: str, **kwargs) -> ProcessingResult:
        """Process DOCX files with formatting preservation."""
        if not DOCX_AVAILABLE:
            return ProcessingResult(
                content_type=ContentType.DOCX,
                extracted_text="",
                metadata={},
                processing_time=0.0,
                confidence_score=0.0,
                error_message="DOCX processing not available"
            )
        
        try:
            doc = Document(file_path)
            
            # Extract text content
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)
            
            extracted_text = '\n\n'.join(text_content)
            
            # Build metadata
            metadata = DocumentMetadata()
            metadata.word_count = len(extracted_text.split())
            metadata.file_size = os.path.getsize(file_path)
            
            # Extract document properties
            try:
                core_props = doc.core_properties
                metadata.title = core_props.title
                metadata.author = core_props.author
                metadata.creation_date = str(core_props.created) if core_props.created else None
                metadata.modification_date = str(core_props.modified) if core_props.modified else None
            except Exception as e:
                logger.warning(f"Could not extract DOCX properties: {e}")
            
            return ProcessingResult(
                content_type=ContentType.DOCX,
                extracted_text=extracted_text,
                metadata=metadata.__dict__,
                processing_time=0.0,
                confidence_score=0.95 if extracted_text.strip() else 0.2
            )
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return ProcessingResult(
                content_type=ContentType.DOCX,
                extracted_text="",
                metadata={},
                processing_time=0.0,
                confidence_score=0.0,
                error_message=str(e)
            )
    
    def _process_image(self, file_path: str, **kwargs) -> ProcessingResult:
        """Process images with OCR and vision-language models."""
        if not OCR_AVAILABLE:
            return ProcessingResult(
                content_type=ContentType.IMAGE,
                extracted_text="",
                metadata={},
                processing_time=0.0,
                confidence_score=0.0,
                error_message="OCR processing not available"
            )
        
        try:
            # Load image
            image = Image.open(file_path)
            
            # Extract text using OCR
            extracted_text = pytesseract.image_to_string(image)
            
            # Build metadata
            metadata = DocumentMetadata()
            metadata.file_size = os.path.getsize(file_path)
            metadata.word_count = len(extracted_text.split())
            
            # Add image-specific metadata
            image_metadata = {
                'width': image.width,
                'height': image.height,
                'format': image.format,
                'mode': image.mode
            }
            metadata.__dict__.update(image_metadata)
            
            # Enhance with vision-language model if available
            if self.ollama_client and kwargs.get('use_vision', True):
                try:
                    enhanced_text = self._enhance_with_vision_model(file_path, extracted_text)
                    if enhanced_text:
                        extracted_text = f"{extracted_text}\n\n[Vision Model Analysis]\n{enhanced_text}"
                except Exception as e:
                    logger.warning(f"Vision model enhancement failed: {e}")
            
            confidence = 0.8 if extracted_text.strip() else 0.1
            
            return ProcessingResult(
                content_type=ContentType.IMAGE,
                extracted_text=extracted_text,
                metadata=metadata.__dict__,
                processing_time=0.0,
                confidence_score=confidence
            )
            
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            return ProcessingResult(
                content_type=ContentType.IMAGE,
                extracted_text="",
                metadata={},
                processing_time=0.0,
                confidence_score=0.0,
                error_message=str(e)
            )
    
    def _process_text(self, file_path: str, **kwargs) -> ProcessingResult:
        """Process plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
            
            metadata = DocumentMetadata()
            metadata.word_count = len(extracted_text.split())
            metadata.file_size = os.path.getsize(file_path)
            
            return ProcessingResult(
                content_type=ContentType.TEXT,
                extracted_text=extracted_text,
                metadata=metadata.__dict__,
                processing_time=0.0,
                confidence_score=1.0
            )
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            return ProcessingResult(
                content_type=ContentType.TEXT,
                extracted_text="",
                metadata={},
                processing_time=0.0,
                confidence_score=0.0,
                error_message=str(e)
            )
    
    def _process_html(self, file_path: str, **kwargs) -> ProcessingResult:
        """Process HTML files with tag removal."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Simple HTML tag removal (could be enhanced with BeautifulSoup)
            import re
            text_content = re.sub(r'<[^>]+>', '', html_content)
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            metadata = DocumentMetadata()
            metadata.word_count = len(text_content.split())
            metadata.file_size = os.path.getsize(file_path)
            
            return ProcessingResult(
                content_type=ContentType.HTML,
                extracted_text=text_content,
                metadata=metadata.__dict__,
                processing_time=0.0,
                confidence_score=0.85
            )
            
        except Exception as e:
            logger.error(f"Error processing HTML file {file_path}: {e}")
            return ProcessingResult(
                content_type=ContentType.HTML,
                extracted_text="",
                metadata={},
                processing_time=0.0,
                confidence_score=0.0,
                error_message=str(e)
            )
    
    def _process_markdown(self, file_path: str, **kwargs) -> ProcessingResult:
        """Process Markdown files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            # Simple markdown processing (could be enhanced with markdown parser)
            import re
            # Remove markdown formatting
            text_content = re.sub(r'#+\s*', '', markdown_content)  # Headers
            text_content = re.sub(r'\*\*(.*?)\*\*', r'\1', text_content)  # Bold
            text_content = re.sub(r'\*(.*?)\*', r'\1', text_content)  # Italic
            text_content = re.sub(r'`(.*?)`', r'\1', text_content)  # Code
            text_content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text_content)  # Links
            
            metadata = DocumentMetadata()
            metadata.word_count = len(text_content.split())
            metadata.file_size = os.path.getsize(file_path)
            
            return ProcessingResult(
                content_type=ContentType.MARKDOWN,
                extracted_text=text_content,
                metadata=metadata.__dict__,
                processing_time=0.0,
                confidence_score=0.9
            )
            
        except Exception as e:
            logger.error(f"Error processing Markdown file {file_path}: {e}")
            return ProcessingResult(
                content_type=ContentType.MARKDOWN,
                extracted_text="",
                metadata={},
                processing_time=0.0,
                confidence_score=0.0,
                error_message=str(e)
            )
    
    def _enhance_with_vision_model(self, image_path: str, ocr_text: str) -> Optional[str]:
        """Enhance image understanding with vision-language models via Ollama."""
        if not self.ollama_client:
            return None
        
        try:
            # Check for vision-capable models
            models = self.ollama_client.list()
            vision_models = ['llava', 'bakllava', 'llava-phi3']  # Common vision models
            
            available_vision_model = None
            for model in models.get('models', []):
                model_name = model.get('name', '').lower()
                if any(vm in model_name for vm in vision_models):
                    available_vision_model = model.get('name')
                    break
            
            if not available_vision_model:
                logger.info("No vision-capable models found in Ollama")
                return None
            
            # Encode image to base64
            with open(image_path, 'rb') as img_file:
                img_data = base64.b64encode(img_file.read()).decode('utf-8')
            
            # Create prompt for vision model
            prompt = f"""Analyze this image and provide insights beyond what OCR can detect:

OCR Text Found: {ocr_text[:500]}...

Please describe:
1. Visual elements not captured by OCR
2. Context and meaning of the content
3. Document structure and layout
4. Any charts, diagrams, or visual information

Provide a concise analysis:"""

            # Call vision model
            response = self.ollama_client.generate(
                model=available_vision_model,
                prompt=prompt,
                images=[img_data],
                options={'temperature': 0.3}
            )
            
            return response.get('response', '')
            
        except Exception as e:
            logger.warning(f"Vision model enhancement failed: {e}")
            return None
    
    def _enhance_with_hierarchical_processing(self, result: ProcessingResult, **kwargs) -> ProcessingResult:
        """Enhance results with hierarchical summarization."""
        try:
            if not result.extracted_text or len(result.extracted_text.strip()) < 50:
                return result
            
            # Apply hierarchical processing
            hierarchical_result = self.hierarchical_engine.process_text(
                result.extracted_text,
                kwargs.get('hierarchical_config', {})
            )
            
            # Add hierarchical analysis to metadata
            result.metadata['hierarchical_analysis'] = hierarchical_result
            
            # Enhance confidence based on successful processing
            if hierarchical_result and 'error' not in hierarchical_result:
                result.confidence_score = min(result.confidence_score + 0.1, 1.0)
            
            return result
            
        except Exception as e:
            logger.warning(f"Hierarchical enhancement failed: {e}")
            return result
    
    def process_batch(self, file_paths: List[str], **kwargs) -> List[ProcessingResult]:
        """Process multiple files in batch."""
        results = []
        
        for file_path in file_paths:
            try:
                result = self.process_file(file_path, **kwargs)
                results.append(result)
                
                logger.info(f"Processed {file_path}: {result.content_type.value}, "
                           f"confidence: {result.confidence_score:.2f}")
                
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                results.append(ProcessingResult(
                    content_type=ContentType.UNKNOWN,
                    extracted_text="",
                    metadata={},
                    processing_time=0.0,
                    confidence_score=0.0,
                    error_message=str(e)
                ))
        
        return results
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self._supported_formats.keys())
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get processing statistics and capabilities."""
        return {
            'supported_formats': len(self._supported_formats),
            'pdf_available': PDF_AVAILABLE,
            'docx_available': DOCX_AVAILABLE,
            'ocr_available': OCR_AVAILABLE,
            'vision_available': VISION_AVAILABLE,
            'ollama_available': self.ollama_client is not None,
            'hierarchical_engine': True
        }


# Example usage and testing
if __name__ == "__main__":
    # Initialize processor
    processor = MultiModalProcessor()
    
    # Print capabilities
    stats = processor.get_processing_stats()
    print("Multi-Modal Processor Capabilities:")
    for key, value in stats.items():
        print(f"  {key}: {value}")
    
    print(f"\nSupported formats: {processor.get_supported_formats()}")